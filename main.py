import io
import os
import re
import tempfile
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from paddleocr import PaddleOCR, PPStructureV3
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT

app = FastAPI()

# Holds the most recently extracted result in memory (NOT saved to disk).
# /download reads from this instead of re-running OCR + structure detection.
latest_extraction = None

ocr = PaddleOCR(use_textline_orientation=True,
                 enable_mkldnn=False,
                 use_doc_orientation_classify=True,
                 use_doc_unwarping=True,
                 lang="en")

structure_pipeline = PPStructureV3(
    lang="en",
    use_table_recognition=True,
    use_formula_recognition=False,
    use_chart_recognition=False,
    use_seal_recognition=False,
)

ALLOWED_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tiff", ".pdf"}


def parse_html_table(html):
    """Converts PP-StructureV3's HTML table output into a list of rows,
    each row a list of cell strings — the format python-docx needs."""
    rows = []
    for row_match in re.finditer(r"<tr[^>]*>(.*?)</tr>", html, re.S | re.I):
        row_html = row_match.group(1)
        cells = []
        for cell_match in re.finditer(r"<t[dh][^>]*>(.*?)</t[dh]>", row_html, re.S | re.I):
            cell_text = re.sub(r"<[^>]+>", " ", cell_match.group(1))
            cell_text = re.sub(r"\s+", " ", cell_text).strip()
            cells.append(cell_text)
        if cells:
            rows.append(cells)
    return rows


def box_center(poly):
    xs = [p[0] for p in poly]
    ys = [p[1] for p in poly]
    return sum(xs) / len(xs), sum(ys) / len(ys)


def group_into_rows(lines_with_poly, y_threshold=12):
    """Groups OCR lines that sit at roughly the same height into one visual
    row (e.g. 'BILLED TO:' and 'Really Great Company' sitting side by side),
    so they become one tab-separated line instead of two stacked lines.
    Input must already be sorted by (y, x). Returns list of rows, each row a
    list of texts left-to-right."""
    rows = []
    current_row = []
    current_y = None

    for poly, text in lines_with_poly:
        cy = box_center(poly)[1]
        if current_y is None or abs(cy - current_y) <= y_threshold:
            current_row.append((poly, text))
            current_y = cy if current_y is None else current_y
        else:
            rows.append(current_row)
            current_row = [(poly, text)]
            current_y = cy

    if current_row:
        rows.append(current_row)

    result = []
    for row in rows:
        row_sorted = sorted(row, key=lambda item: box_center(item[0])[0])
        result.append([text for _, text in row_sorted])
    return result


def get_structured_pages(tmp_path):
    """
    Runs OCR + PP-StructureV3 on the file and returns a list of pages.
    Each page is a list of blocks:
        {"type": "table", "rows": [[cell, cell, ...], ...]}
        {"type": "text", "paragraphs": ["line one", "label\tvalue", ...]}
    A paragraph string containing "\t" means it was a side-by-side row
    (e.g. label + value) and should get a tab stop when rendered to docx.
    """
    result = ocr.predict(tmp_path)

    page_lines = []  # per page: list of (poly, text)
    for page in result:
        texts = page.get("rec_texts", [])
        polys = page.get("rec_polys", [])
        page_lines.append([
            (poly.tolist() if hasattr(poly, "tolist") else poly, text)
            for text, poly in zip(texts, polys)
        ])

    structure_result = structure_pipeline.predict(tmp_path)

    pages = []

    for page_idx, page_struct in enumerate(structure_result):
        blocks = page_struct.get("parsing_res_list", [])
        blocks = sorted(
            blocks,
            key=lambda b: (
                (getattr(b, "bbox", None) or [0, 0, 0, 0])[1],
                (getattr(b, "bbox", None) or [0, 0, 0, 0])[0],
            ),
        )

        lines_for_page = page_lines[page_idx] if page_idx < len(page_lines) else []
        page_blocks = []

        for block in blocks:
            label = getattr(block, "label", None)
            content = getattr(block, "content", "")
            bbox = getattr(block, "bbox", None)

            if label == "table":
                rows = parse_html_table(content)
                if rows:
                    page_blocks.append({"type": "table", "rows": rows})
                continue

            if bbox:
                x1, y1, x2, y2 = bbox
                lines_in_block = [
                    (poly, text) for poly, text in lines_for_page
                    if x1 - 5 <= box_center(poly)[0] <= x2 + 5
                    and y1 - 5 <= box_center(poly)[1] <= y2 + 5
                ]
                lines_in_block.sort(key=lambda item: (box_center(item[0])[1], box_center(item[0])[0]))
                row_groups = group_into_rows(lines_in_block)
            else:
                row_groups = [[l] for l in content.splitlines() if l.strip()]

            # Single-item rows are just wrapped lines of the same paragraph —
            # merge with spaces into one flowing paragraph. Multi-item rows
            # (side-by-side, e.g. label + value) stay on their own line.
            paragraphs = []
            buffer_lines = []

            def flush_buffer():
                if buffer_lines:
                    paragraphs.append(" ".join(buffer_lines))
                    buffer_lines.clear()

            for row_texts in row_groups:
                if not any(t.strip() for t in row_texts):
                    continue
                if len(row_texts) == 1:
                    buffer_lines.append(row_texts[0])
                else:
                    flush_buffer()
                    paragraphs.append("\t".join(row_texts))

            flush_buffer()

            if paragraphs:
                page_blocks.append({"type": "text", "paragraphs": paragraphs})

        pages.append(page_blocks)

    return pages


def render_preview_text(pages):
    """Flattens the structured pages into plain text for the JSON preview —
    this mirrors exactly what will end up in the Word doc."""
    page_texts = []
    for page in pages:
        lines = []
        for block in page:
            if block["type"] == "table":
                for row in block["rows"]:
                    lines.append(" | ".join(row))
            else:
                lines.extend(block["paragraphs"])
        page_texts.append("\n".join(lines))
    return "\n\n--- Page Break ---\n\n".join(page_texts)


def build_docx(pages):
    """Builds a Word document entirely in memory — nothing is saved to disk."""
    doc = Document()

    for page_idx, page in enumerate(pages):
        for block in page:
            if block["type"] == "table":
                rows = block["rows"]
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j, cell_text in enumerate(row):
                        table.cell(i, j).text = cell_text
                doc.add_paragraph("")
            else:
                for ptext in block["paragraphs"]:
                    p = doc.add_paragraph(ptext)
                    if "\t" in ptext:
                        p.paragraph_format.tab_stops.add_tab_stop(Inches(2.2), WD_TAB_ALIGNMENT.LEFT)

        if page_idx != len(pages) - 1:
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.get("/")
async def home():
    return FileResponse("static/index.html")


@app.post("/ocr")
async def extract_text(file: UploadFile = File(...)):
    global latest_extraction

    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    try:
        pages = get_structured_pages(tmp_path)
        full_text = render_preview_text(pages)

        latest_extraction = {"pages": pages, "filename": file.filename}

        return {
            "filename": file.filename,
            "num_pages": len(pages),
            "text": full_text,
        }
    finally:
        os.remove(tmp_path)


@app.get("/download")
async def download_docx():
    if latest_extraction is None:
        raise HTTPException(status_code=404, detail="No extraction available yet. Please extract text first.")

    buffer = build_docx(latest_extraction["pages"])
    out_name = os.path.splitext(latest_extraction["filename"])[0] + "_extracted.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'},
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="localhost", port=8000)